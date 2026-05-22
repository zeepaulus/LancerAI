"""Document Worker: CV export to PDF/DOCX using professional ATS templates.

Runs as Celery async task to avoid blocking the API.
"""

import asyncio
import base64
import io
import logging
from typing import Any

from celery import shared_task  # type: ignore[import-untyped]

from app.core.providers.connectors import get_llm_connector
from app.service.optimization.template_renderer import CVTemplateRenderer

# Import the Celery app so @shared_task can find and register with it
from app.workers.celery_app import celery_app as _celery_app  # noqa: F401

logger = logging.getLogger(__name__)


def _run_async(coro: Any) -> Any:
    """Helper to run async coroutines in a synchronous context."""
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop = asyncio.get_event_loop()

    if loop.is_running():
        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(asyncio.run, coro)
            return future.result()
    else:
        return loop.run_until_complete(coro)


def _generate_docx_bytes(cv_data: dict[str, Any]) -> bytes:
    """Generate a clean, professional DOCX layout using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Style default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Personal Information
    pi = cv_data.get("personal_info", {})
    name = pi.get("name", "Ứng viên")
    email = pi.get("email", "")
    phone = pi.get("phone", "")
    linkedin = pi.get("linkedin", "")
    location = pi.get("location", "")

    # Add Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(name)
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0x2D, 0x3E, 0x50)
    p_name.paragraph_format.space_after = Pt(2)

    # Add Contact info
    contact_parts = [p for p in [email, phone, linkedin, location] if p]
    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact = p_contact.add_run("  |  ".join(contact_parts))
        run_contact.font.size = Pt(9.5)
        p_contact.paragraph_format.space_after = Pt(18)

    # Helper to add section headings
    def add_heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x3E, 0x50)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True

    # Education
    education_list = cv_data.get("education", [])
    if education_list:
        add_heading("HỌC VẤN")
        for edu in education_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run_school = p.add_run(edu.get("school", ""))
            run_school.bold = True

            period = edu.get("period", "")
            if period:
                run_period = p.add_run(f"\t\t({period})")
                run_period.italic = True

            p_details = doc.add_paragraph()
            p_details.paragraph_format.space_after = Pt(8)
            deg_major = f"{edu.get('degree', '')} — {edu.get('major', '')}"
            gpa = edu.get("gpa", "")
            if gpa:
                deg_major += f" (GPA: {gpa})"
            p_details.add_run(deg_major)

    # Experience
    experience_list = cv_data.get("experience", [])
    if experience_list:
        add_heading("KINH NGHIỆM LÀM VIỆC")
        for exp in experience_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            role_company = f"{exp.get('title', '')} — {exp.get('company', '')}"
            run_rc = p.add_run(role_company)
            run_rc.bold = True

            period = exp.get("period", "")
            if period:
                run_period = p.add_run(f"\t\t({period})")
                run_period.italic = True

            impacts = exp.get("key_impacts", []) or exp.get("descriptions", [])
            if impacts:
                for imp in impacts:
                    p_imp = doc.add_paragraph(style="List Bullet")
                    p_imp.paragraph_format.space_after = Pt(2)
                    p_imp.add_run(str(imp))

            # Add spacing at the end of each entry
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(2)
            p_space.paragraph_format.space_after = Pt(4)

    # Projects
    projects_list = cv_data.get("projects", [])
    if projects_list:
        add_heading("DỰ ÁN")
        for proj in projects_list:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run_proj = p.add_run(proj.get("name", ""))
            run_proj.bold = True

            role = proj.get("role", "")
            if role:
                p.add_run(f" — {role}").italic = True

            desc = proj.get("description", "")
            if desc:
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.space_after = Pt(2)
                p_desc.add_run(desc)

            tech = proj.get("tech_stack", [])
            if tech:
                p_tech = doc.add_paragraph()
                p_tech.paragraph_format.space_after = Pt(8)
                run_tech_lbl = p_tech.add_run("Công nghệ sử dụng: ")
                run_tech_lbl.italic = True
                p_tech.add_run(", ".join(tech))

    # Skills Matrix
    skills = cv_data.get("skills_matrix", {})
    if skills:
        add_heading("KỸ NĂNG")
        for cat, label in [
            ("languages", "Ngôn ngữ lập trình"),
            ("frameworks", "Frameworks & Thư viện"),
            ("tools", "Công cụ & Hệ điều hành"),
            ("soft_skills", "Kỹ năng mềm"),
        ]:
            items = skills.get(cat, [])
            if items:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"{label}: ").bold = True
                p.add_run(", ".join(items))

    # Certifications
    certs = cv_data.get("certifications", [])
    if certs:
        add_heading("CHỨNG CHỈ")
        for cert in certs:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.add_run(str(cert))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()


@shared_task(bind=True, max_retries=2, default_retry_delay=30)  # type: ignore[untyped-decorator]
def generate_document(
    self: Any,
    cv_data: dict[str, Any],
    template: str = "standard_ats",
    output_format: str = "pdf",
) -> dict[str, Any]:
    """Generate a professional CV document from structured data.

    Templates: standard_ats, modern_tech, management.
    Formats: pdf (via WeasyPrint), docx (via python-docx).
    """
    logger.info("Starting CV export. format=%s, template=%s", output_format, template)

    if output_format.lower() == "pdf":
        template_map = {
            "standard_ats": "harvard",
            "modern_tech": "modern",
            "management": "creative",
        }
        mapped_template = template_map.get(template, template)
        if mapped_template not in {"harvard", "modern", "minimal", "creative"}:
            mapped_template = "harvard"

        try:
            llm = get_llm_connector()
            renderer = CVTemplateRenderer(llm)
            # Use async renderer inside helper
            pdf_bytes = _run_async(renderer.render_pdf(cv_data, template=mapped_template))
            doc_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
            return {
                "status": "success",
                "template": template,
                "output_format": "pdf",
                "document_b64": doc_b64,
            }
        except Exception as exc:
            logger.error("PDF generation failed: %s", exc, exc_info=True)
            try:
                self.retry(exc=exc)
            except Exception:
                raise exc from None
            return {}

    elif output_format.lower() == "docx":
        try:
            docx_bytes = _generate_docx_bytes(cv_data)
            doc_b64 = base64.b64encode(docx_bytes).decode("utf-8")
            return {
                "status": "success",
                "template": template,
                "output_format": "docx",
                "document_b64": doc_b64,
            }
        except Exception as exc:
            logger.error("DOCX generation failed: %s", exc, exc_info=True)
            try:
                self.retry(exc=exc)
            except Exception:
                raise exc from None
            return {}

    else:
        raise ValueError(f"Unsupported output format: {output_format}")
